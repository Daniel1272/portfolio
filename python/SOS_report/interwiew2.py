import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from io import BytesIO
import re


df = pd.read_excel('Frontline Volunteers Daily Return Form- test data for DAA.xlsx')
pd.set_option('display.max_rows', None)
pd.set_option('display.max_columns', None)


# preprocessing
print(df.info())


# delete last 5 lines as they are empty
df = df.iloc[0:5]
df: pd.DataFrame


# Total number of calls answered.
total_calls_answered = df[('How many calls did you take today?  This is the TOTAL number of calls '
                           '(it does not include hang up calls- these go below), thank you')].sum()
print('Total number of calls answered: ', total_calls_answered)


# Total number of calls answered broken down by volunteer.
calls_answered_by_volunteer = df[['What is your name and alias',
                                  'How many calls did you take today?  This is the TOTAL number of calls '
                                  '(it does not include hang up calls- these go below), thank you']]
print('Total number of calls answered broken down by volunteer: ', calls_answered_by_volunteer, '', sep='\n')


# Total number of calls with Female callers.
female_callers = df['How many callers identified as female?'].sum()
print('Total number of calls with Female callers: ', female_callers)


# Number of calls with Female callers broken down by volunteer.
female_calls_answered_by_volunteer = df[['What is your name and alias', 'How many callers identified as female?']]
print('Number of calls with Female callers broken down by volunteer: ', female_calls_answered_by_volunteer, '', sep='\n')


# Total number of calls with Male callers.
male_callers = df['How many callers identified as male?'].sum()
print('Total number of calls with Male callers: ', male_callers)


# Number of calls with Female callers broken down by volunteer.
male_calls_answered_by_volunteer = df[['What is your name and alias', 'How many callers identified as male?']]
print('Number of calls with Male callers broken down by volunteer: ', male_calls_answered_by_volunteer, '', sep='\n')


# Total number of Hang ups calls.
hang_ups_calls = df['How many hang up calls today?'].sum()
print('Total number of Hang ups calls: ', hang_ups_calls)

# Number of Hang ups calls broken down by volunteer.
hang_ups_calls_by_volunteer = df[['What is your name and alias', 'How many hang up calls today?']]
print('Number of Hang ups calls broken down by volunteer: ', hang_ups_calls_by_volunteer, '', sep='\n')


# Total number of Immediate Risk (IR) calls.
ir_calls = df['How many immediate risk calls did you have today?'].sum()
print('How many immediate risk calls did you have today?: ', ir_calls)

# Which volunteer(s) had IR calls.
ir_calls_by_volunteer = df[['What is your name and alias', 'How many immediate risk calls did you have today?',
                            'Please tell us all about your IR calls, including any action taken by you or your'
                            ' spot checker.  Remember to include caller\'s name, telephone number & reason '
                            'for IR.  You should use the following ...']][df[('How many immediate risk calls '
                                                                              'did you have today?')] > 0]
print(ir_calls_by_volunteer)


# Total number of Minor calls
minor_calls = df['Today, did you have any calls from minors (under 18)?'][df[('Today, did you have any calls from minors (under 18)?')] == 'Yes'].count()
print('Total number of Minor calls: ',minor_calls)


# Which volunteer(s) had Minor calls.
minor_calls_by_volunteer = df[['What is your name and alias',
                               'Tell us about these calls including names, ages, phone numbers,'
                               ' call history, outcome, actions taken and if you completed a MASH'
                               ' form or not.  From now on, you should send a copy of your completed...',
                               'Did you offer to, or did the minor ask you, to speak to the parents at any '
                               'point during the call?','If you spoke to parents, please give us further'
                               ' information here.  Otherwise type NA','Were any minors already under, '
                               'or seeking the care and/or support of:',
                               'What were the reasons for the minors calling today?',

                               ]][df[(''
                 'Today, did you have any calls from minors (under 18)?')] == 'Yes']

print(minor_calls_by_volunteer)


# Total number of Abusive calls.
abusive_calls = df['Did you have callers who were abusive, threatening, sexual or beyond your boundaries?'].sum()
print('Total number of Abusive calls: ', abusive_calls)

# Which volunteer had Abusive calls?
abusive_calls_by_volunteer = df[['What is your name and alias','Tell us about the callers'
                                 ' in the question above otherwise type NA',]][df[('Did you'
                                 ' have callers who were abusive, threatening, sexual or beyond'
                                 ' your boundaries?')]>0]
print('Which volunteer had Abusive calls?:', abusive_calls_by_volunteer)


# Were there any Nimvelo issues? If yes, what were they and which volunteers had Nimvelo issues?

nimvelo_issues = df[['What is your name and alias','If there were any Nimvelo issues please '
                     'let us know what there were. If there were not any Nimvelo issues please type NA',
                     ]][df['Were there any Nimvelo issues during your shift?'] == 'Yes']
print('Nimvelo issues: ',nimvelo_issues,'',sep='\n')


# How are the volunteers performing? e.g., are they submitting their return forms within 48
# hours of their shift?
volunteers_performing = df[['What is your name and alias']].copy()
volunteers_performing['submitting_their_return_forms'] = (df.loc[:,'What date are you completing this form']
                                                          - df.loc[:,'What date does your shift refer to'])

print('How are the volunteers submitting their return forms: ',volunteers_performing)


# The welfare of our volunteers including
volunteers_impacted = df[['What is your name and alias','Tell us about any impacts directly caused'
                          ' by volunteering for SOS','How well supported do you feel by the SOS team generally?',
                          'Are you enjoying your time volunteering with SOS?']]

print('The welfare of our volunteers: ',volunteers_impacted)


# Create a new Document
doc = Document()

# Title
doc.add_heading('Volunteer Call Statistics', 0)

# Add Total Number of Calls Answered
doc.add_paragraph('Total number of calls answered: ', style='Heading 1').bold = True
doc.add_paragraph(f'{total_calls_answered}')

# Add Calls Answered by Volunteer
doc.add_paragraph('Total number of calls answered broken down by volunteer:', style='Heading 1').bold = True
for index, row in calls_answered_by_volunteer.iterrows():
    doc.add_paragraph(f"{row['What is your name and alias']}:"
                      f"{row[('How many calls did you take today?  This is the TOTAL number of calls (it does not include hang up calls- these go below), thank you')]}")

# Add Total Number of Female Callers
doc.add_paragraph('Total number of calls with Female callers: ', style='Heading 1').bold = True
doc.add_paragraph(f'{female_callers}')

# Add Female Calls Answered by Volunteer
doc.add_paragraph('Number of calls with Female callers broken down by volunteer:', style='Heading 1').bold = True
for index, row in female_calls_answered_by_volunteer.iterrows():
    doc.add_paragraph(f"{row['What is your name and alias']}: {row['How many callers identified as female?']}")

# Add Total Number of Male Callers
doc.add_paragraph('Total number of calls with Male callers: ', style='Heading 1').bold = True
doc.add_paragraph(f'{male_callers}')

# Add Male Calls Answered by Volunteer
doc.add_paragraph('Number of calls with Male callers broken down by volunteer:', style='Heading 1').bold = True
for index, row in male_calls_answered_by_volunteer.iterrows():
    doc.add_paragraph(f"{row['What is your name and alias']}: {row['How many callers identified as male?']}")

# Add Total Number of Hang-ups
doc.add_paragraph('Total number of Hang ups calls: ', style='Heading 1').bold = True
doc.add_paragraph(f'{hang_ups_calls}')

# Add Hang-ups Calls by Volunteer
doc.add_paragraph('Number of Hang ups calls broken down by volunteer:', style='Heading 1').bold = True
for index, row in hang_ups_calls_by_volunteer.iterrows():
    doc.add_paragraph(f"{row['What is your name and alias']}: {row['How many hang up calls today?']}")

# Add Total Immediate Risk Calls
doc.add_paragraph('Total number of Immediate Risk (IR) calls: ', style='Heading 1').bold = True
doc.add_paragraph(f'{ir_calls}')

# Add IR Calls by Volunteer
doc.add_paragraph('Which volunteer(s) had IR calls:', style='Heading 1').bold = True
for index, row in ir_calls_by_volunteer.iterrows():
    cleaned_text = re.sub(r'\u00A0', ' ', row[
        'Please tell us all about your IR calls, including any action taken by you or your'
                            ' spot checker.  Remember to include caller\'s name, telephone number & reason '
                            'for IR.  You should use the following ...'])

    doc.add_paragraph(f"{row['What is your name and alias']}:"
                      f" {row['How many immediate risk calls did you have today?']} {cleaned_text}")

# Add Total Minor Calls
doc.add_paragraph('Total number of Minor calls: ', style='Heading 1').bold = True
doc.add_paragraph(f'{minor_calls}')

# Add Minor Calls by Volunteer
doc.add_paragraph('Which volunteer(s) had Minor calls:', style='Heading 1').bold = True
for index, row in minor_calls_by_volunteer.iterrows():
    doc.add_paragraph(f"{row['What is your name and alias']}: {row['Tell us about these calls including names, ages, phone numbers, call history, outcome, actions taken and if you completed a MASH form or not.  From now on, you should send a copy of your completed...']}")

# Add Total Abusive Calls
doc.add_paragraph('Total number of Abusive calls: ', style='Heading 1').bold = True
doc.add_paragraph(f'{abusive_calls}')

# Add Abusive Calls by Volunteer
doc.add_paragraph('Which volunteer had Abusive calls:', style='Heading 1').bold = True
for index, row in abusive_calls_by_volunteer.iterrows():
    doc.add_paragraph(f"{row['What is your name and alias']}: {row['Tell us about the callers in the question above otherwise type NA']}")

# Add Nimvelo Issues
doc.add_paragraph(f'Were there any Nimvelo issues?', style='Heading 1').bold = True
for index, row in nimvelo_issues.iterrows():
    doc.add_paragraph(f"{row['What is your name and alias']}: {row['If there were any Nimvelo issues please let us know what there were. If there were not any Nimvelo issues please type NA']}")

# Add Volunteer Performance
doc.add_paragraph(f'Volunteer performance:', style='Heading 1').bold = True
for index, row in volunteers_performing.iterrows():
    doc.add_paragraph(f"{row['What is your name and alias']}: {'Submitted in time' if row['submitting_their_return_forms'].days <= 2 else 'Late Submission'} "
                      f"({row['submitting_their_return_forms'].days} days)")

# Add Welfare Information
doc.add_paragraph(f'The welfare of our volunteers:', style='Heading 1').bold = True
for index, row in volunteers_impacted.iterrows():
    doc.add_paragraph(f"{row['What is your name and alias']}: {row['Tell us about any impacts directly caused by volunteering for SOS']}")

# Visualization 1: Calls Answered by Volunteer (Bar Chart)
fig, ax = plt.subplots()
calls_per_volunteer = df.groupby('What is your name and alias')['How many calls did you take today?  This is the TOTAL number of calls (it does not include hang up calls- these go below), thank you'].sum()
calls_per_volunteer.plot(kind='bar', ax=ax)
ax.set_title('Calls Answered by Volunteer')
ax.set_ylabel('Number of Calls')
plt.tight_layout()

# Save the plot to a BytesIO object
img_stream = BytesIO()
plt.savefig(img_stream, format='png')
img_stream.seek(0)

# Add the plot to the Word document
doc.add_paragraph('Visualization: Calls Answered by Volunteer', style='Heading 1').bold = True
doc.add_picture(img_stream)

# Visualization 2: Gender Distribution of Callers (Pie Chart)
fig2, ax2 = plt.subplots()
gender_counts = df[['How many callers identified as female?', 'How many callers identified as male?']].sum()
gender_counts.plot(kind='pie', ax=ax2, autopct='%1.1f%%', startangle=90, labels=['Female', 'Male'])
ax2.set_ylabel('')
plt.tight_layout()

# Save the pie chart to a BytesIO object
img_stream2 = BytesIO()
plt.savefig(img_stream2, format='png')
img_stream2.seek(0)

# Add the pie chart to the Word document
doc.add_paragraph('Visualization: Gender Distribution of Callers', style='Heading 1').bold = True
doc.add_picture(img_stream2)

# Visualization 3: Hang-up Calls (Bar Chart)
fig3, ax3 = plt.subplots()
hang_ups_per_volunteer = df.groupby('What is your name and alias')['How many hang up calls today?'].sum()
hang_ups_per_volunteer.plot(kind='bar', ax=ax3, color='red')
ax3.set_title('Hang-up Calls by Volunteer')
ax3.set_ylabel('Number of Hang-ups')
plt.tight_layout()

# Save the plot to a BytesIO object
img_stream3 = BytesIO()
plt.savefig(img_stream3, format='png')
img_stream3.seek(0)

# Add the hang-up plot to the Word document
doc.add_paragraph('Visualization: Hang-up Calls by Volunteer', style='Heading 1').bold = True
doc.add_picture(img_stream3)

# Save the document
doc.save('Volunteer_Call_Report_with_Visualizations.docx')

print("Document with Visualizations created successfully!")
